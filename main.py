import customtkinter as ctk
from database import DBManager
from gui_components import AppointmentModal, ServiceModal, MasterModal, ClientModal, LoginModal, RegisterModal, COLORS
from tkinter import messagebox, filedialog
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import os

class SilkWayApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        self.title("КРАСНАЯ КОРОЛЕВА")
        self.geometry("1450x900")
        ctk.set_appearance_mode("dark")
        
        self.db = DBManager()
        self.current_user = None

        self.nav_frame = ctk.CTkFrame(self, width=280, corner_radius=0, fg_color="#0a0a0a")
        self.nav_frame.pack_propagate(False)

        self.main_view = ctk.CTkFrame(self, fg_color="#121212", corner_radius=0)
        
        self.journal_data = []
        self.clients_data = []
        
        self.after(100, self.require_login)

    def require_login(self):
        self.nav_frame.pack_forget()
        self.main_view.pack_forget()
        LoginModal(self, self.db, self.on_login_success)

    def on_login_success(self, user):
        self.current_user = user
        self.nav_frame.pack(side="left", fill="y")
        self.main_view.pack(side="right", fill="both", expand=True)
        self.build_ui()

    def build_ui(self):
        for widget in self.nav_frame.winfo_children():
            widget.destroy()

        ctk.CTkLabel(self.nav_frame, text="•────•♛•────•\nКРАСНАЯ\nКОРОЛЕВА", font=("Arial", 28, "bold"), text_color="#ff4757").pack(pady=(40, 5))
        role_txt = "АДМИНИСТРАТОР" if self.current_user['role'] == 'admin' else "МАСТЕР"
        ctk.CTkLabel(self.nav_frame, text=role_txt, font=("Arial", 11, "bold"), text_color="#2ecc71").pack(pady=(0, 60))

        if self.current_user['role'] == 'admin':
            self.create_nav_btn("📋  Журнал записей", self.show_journal)
            self.create_nav_btn("👥  Клиенты", self.show_clients)
            self.create_nav_btn("💰  Прайс-лист", self.show_services)
            self.create_nav_btn("👤  Мастера", self.show_masters)
            self.create_nav_btn("📊  Аналитика", self.show_analytics)
            self.create_nav_btn("🔐  Доступ", self.show_registration)
            
            ctk.CTkButton(self.nav_frame, text="ВЫЙТИ ИЗ СИСТЕМЫ", fg_color="#c0392b", hover_color="#922b21", height=45,
                          command=self.logout).pack(side="bottom", pady=20, padx=30, fill="x")

            ctk.CTkButton(self.nav_frame, text="+ НОВАЯ ЗАПИСЬ", fg_color="#2ecc71", hover_color="#27ae60", 
                          height=55, font=("Arial", 15, "bold"), corner_radius=12,
                          command=lambda: AppointmentModal(self, self.db, self.show_journal)).pack(side="bottom", pady=10, padx=30, fill="x")
        else:
            self.create_nav_btn("📋  Мои записи", self.show_journal)
            ctk.CTkButton(self.nav_frame, text="ВЫЙТИ ИЗ СИСТЕМЫ", fg_color="#c0392b", hover_color="#922b21", height=45,
                          command=self.logout).pack(side="bottom", pady=50, padx=30, fill="x")

        self.show_journal()

    def logout(self):
        if messagebox.askyesno("Выход", "Вы действительно хотите сменить пользователя?"):
            self.current_user = None
            self.clear_view()
            self.require_login()

    def create_nav_btn(self, text, command):
        btn = ctk.CTkButton(self.nav_frame, text=text, anchor="w", fg_color="transparent", 
                            text_color="#aaa", hover_color="#161616", height=60, 
                            font=("Arial", 16), command=command)
        btn.pack(fill="x", padx=15, pady=3)

    def clear_view(self):
        for widget in self.main_view.winfo_children():
            widget.destroy()

    def show_registration(self):
        RegisterModal(self, self.db)

    def show_journal(self):
        self.clear_view()
        hdr = ctk.CTkFrame(self.main_view, fg_color="transparent")
        hdr.pack(fill="x", padx=60, pady=(50, 30))
        title_text = "♛ РАСПИСАНИЕ" if self.current_user['role'] == 'admin' else "👑 МОИ ВИЗИТЫ"
        ctk.CTkLabel(hdr, text=title_text, font=("Arial", 38, "bold")).pack(side="left")
        
        self.search_var = ctk.StringVar()
        self.search_var.trace_add("write", self.render_journal_list)
        ctk.CTkEntry(hdr, textvariable=self.search_var, placeholder_text="Поиск по имени или услуге...", width=300, height=45).pack(side="right")

        self.journal_container = ctk.CTkScrollableFrame(self.main_view, fg_color="transparent")
        self.journal_container.pack(fill="both", expand=True, padx=40)

        master_id = self.current_user['master_id'] if self.current_user['role'] == 'master' else None
        self.journal_data = self.db.get_journal(master_id=master_id)
        self.render_journal_list()

    def print_receipt(self, app_data):
        try:
            filename = f"Receipt_{app_data['id']}.txt"
            with open(filename, "w", encoding="utf-8") as f:
                f.write("      КРАСНАЯ КОРОЛЕВА      \n")
                f.write("----------------------------------\n")
                f.write(f"Чек №: {app_data['id']}\n")
                f.write(f"Дата: {app_data['appointment_date'].strftime('%d.%m.%Y %H:%M')}\n")
                f.write(f"Клиент: {app_data['client']}\n")
                f.write(f"Мастер: {app_data['master']}\n")
                f.write("----------------------------------\n")
                f.write(f"Услуги: {app_data['service']}\n")
                f.write(f"ИТОГО К ОПЛАТЕ: {int(app_data['price'])} руб.\n")
                f.write("----------------------------------\n")
                f.write("   Спасибо, что выбираете нас!    \n")
            os.startfile(filename) 
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать чек: {e}")

    def safe_status_update(self, aid, val):
        # Баг исправлен: даем 500мс для закрытия меню выбора, чтобы окно не сворачивалось
        self.db.update_status(aid, val)
        self.after(500, self.show_journal)

    def render_journal_list(self, *args):
        for widget in self.journal_container.winfo_children():
            widget.destroy()

        query = self.search_var.get().lower()
        filtered_data = [r for r in self.journal_data if query in r['client'].lower() or query in r['service'].lower()]

        for r in filtered_data:
            clr = COLORS.get(r['category'])
            dt = r['appointment_date']
            end_dt = dt + timedelta(minutes=int(r['total_duration']))
            
            card = ctk.CTkFrame(self.journal_container, fg_color="#1a1a1a", height=125, corner_radius=15)
            card.pack(fill="x", pady=8); card.pack_propagate(False)
            ctk.CTkFrame(card, width=7, fg_color=clr).pack(side="left", fill="y")

            t_f = ctk.CTkFrame(card, fg_color="transparent")
            t_f.pack(side="left", padx=30)
            ctk.CTkLabel(t_f, text=f"{dt.strftime('%H:%M')} - {end_dt.strftime('%H:%M')}", font=("Arial", 22, "bold")).pack()
            ctk.CTkLabel(t_f, text=dt.strftime('%d.%m'), font=("Arial", 13), text_color="#555").pack()

            info = ctk.CTkFrame(card, fg_color="transparent")
            info.pack(side="left", padx=20, fill="both", expand=True, pady=15)
            
            ctk.CTkLabel(info, text=f"{r['service']} — {int(r['price'])} ₽", font=("Arial", 18, "bold"), text_color=clr, anchor="w").pack(fill="x")
            ctk.CTkLabel(info, text=f"👤 {r['client']}  |  ✂️ {r['master']}", font=("Arial", 15), text_color="#777", anchor="w").pack(fill="x", pady=(5, 0))

            if r['status'] == "Завершено":
                ctk.CTkButton(card, text="🧾 ЧЕК", width=60, height=45, fg_color="#f39c12", hover_color="#e67e22",
                              command=lambda row=r: self.print_receipt(row)).pack(side="right", padx=5)

            if self.current_user['role'] == 'admin':
                ctk.CTkButton(card, text="🗑", width=45, height=45, fg_color="#262626", hover_color="#c0392b",
                              command=lambda aid=r['id']: (self.db.execute("DELETE FROM appointments WHERE id=%s", (aid,)), self.show_journal())).pack(side="right", padx=10)
                ctk.CTkButton(card, text="✏️", width=45, height=45, fg_color="#262626", hover_color="#3498db",
                              command=lambda row=r: AppointmentModal(self, self.db, self.show_journal, data=row)).pack(side="right", padx=5)
                
                s_var = ctk.StringVar(value=r['status'])
                opt = ctk.CTkOptionMenu(card, values=["Ожидание", "В процессе", "Завершено", "Отмена"],
                                        variable=s_var, width=155, height=42, fg_color="#262626",
                                        command=lambda val, aid=r['id']: self.safe_status_update(aid, val))
                opt.pack(side="right", padx=5)
            else:
                vals = [r['status'], "Завершено"] if r['status'] != "Завершено" else ["Завершено"]
                s_var = ctk.StringVar(value=r['status'])
                ctk.CTkOptionMenu(card, values=vals, variable=s_var, width=155, height=42, fg_color="#262626",
                                  command=lambda val, aid=r['id']: self.safe_status_update(aid, val)).pack(side="right", padx=30)

    def show_clients(self):
        self.clear_view()
        hdr = ctk.CTkFrame(self.main_view, fg_color="transparent")
        hdr.pack(fill="x", padx=50, pady=45)
        ctk.CTkLabel(hdr, text="👑 КЛИЕНТСКАЯ БАЗА", font=("Arial", 38, "bold")).pack(side="left")
        
        controls = ctk.CTkFrame(hdr, fg_color="transparent")
        controls.pack(side="right")
        self.client_search_var = ctk.StringVar()
        self.client_search_var.trace_add("write", self.render_clients_list)
        ctk.CTkEntry(controls, textvariable=self.client_search_var, placeholder_text="Поиск...", width=200, height=45).pack(side="left", padx=15)
        ctk.CTkButton(controls, text="+ НОВЫЙ КЛИЕНТ", fg_color="#3498db", height=45, font=("Arial", 14, "bold"),
                      command=lambda: ClientModal(self, self.db, on_save=self.show_clients)).pack(side="left")

        self.clients_container = ctk.CTkScrollableFrame(self.main_view, fg_color="transparent")
        self.clients_container.pack(fill="both", expand=True, padx=40)
        self.clients_data = self.db.fetch("SELECT * FROM clients ORDER BY name")
        self.render_clients_list()

    def render_clients_list(self, *args):
        for widget in self.clients_container.winfo_children():
            widget.destroy()
        query = self.client_search_var.get().lower()
        filtered = [c for c in self.clients_data if query in c['name'].lower() or (c.get('phone') and query in c['phone'])]
        for c in filtered:
            f = ctk.CTkFrame(self.clients_container, fg_color="#1a1a1a", height=100, corner_radius=15)
            f.pack(fill="x", pady=8); f.pack_propagate(False)
            info = ctk.CTkFrame(f, fg_color="transparent")
            info.pack(side="left", padx=30, fill="y")
            ctk.CTkLabel(info, text=c['name'], font=("Arial", 22, "bold"), anchor="w").pack(pady=(15, 0))
            ctk.CTkLabel(info, text=f"📞 {c.get('phone') or 'Нет номера'}", font=("Arial", 14), text_color="#555", anchor="w").pack()
            ctk.CTkButton(f, text="ПРОФИЛЬ", width=140, height=45, fg_color="#262626",
                          command=lambda d=c: ClientModal(self, self.db, d, self.show_clients)).pack(side="right", padx=30)

    def show_services(self):
        self.clear_view()
        hdr = ctk.CTkFrame(self.main_view, fg_color="transparent")
        hdr.pack(fill="x", padx=50, pady=45)
        ctk.CTkLabel(hdr, text="👑 ПРАЙС-ЛИСТ", font=("Arial", 38, "bold")).pack(side="left")
        ctk.CTkButton(hdr, text="+ ДОБАВИТЬ УСЛУГУ", fg_color="#3498db", 
                      command=lambda: ServiceModal(self, self.db, on_save=self.show_services)).pack(side="right")

        scroll = ctk.CTkScrollableFrame(self.main_view, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=40)
        for s in self.db.fetch("SELECT * FROM services ORDER BY category, name"):
            clr = COLORS.get(s['category'])
            f = ctk.CTkFrame(scroll, fg_color="#1a1a1a", height=85, corner_radius=12)
            f.pack(fill="x", pady=6); f.pack_propagate(False)
            ctk.CTkFrame(f, width=7, fg_color=clr).pack(side="left", fill="y")
            ctk.CTkLabel(f, text=s['name'], font=("Arial", 20, "bold"), anchor="w").pack(side="left", padx=30)
            ctk.CTkLabel(f, text=f"{int(s['price'])} ₽", font=("Arial", 22, "bold"), text_color="#2ecc71").pack(side="left", padx=50)
            ctk.CTkButton(f, text="✎", width=50, height=50, fg_color="#262626", 
                          command=lambda d=s: ServiceModal(self, self.db, d, self.show_services)).pack(side="right", padx=25)

    def show_masters(self):
        self.clear_view()
        hdr = ctk.CTkFrame(self.main_view, fg_color="transparent")
        hdr.pack(fill="x", padx=50, pady=45)
        ctk.CTkLabel(hdr, text="👑 НАША КОМАНДА", font=("Arial", 38, "bold")).pack(side="left")
        ctk.CTkButton(hdr, text="+ НОВЫЙ МАСТЕР", fg_color="#3498db", 
                      command=lambda: MasterModal(self, self.db, on_save=self.show_masters)).pack(side="right")
        
        scroll = ctk.CTkScrollableFrame(self.main_view, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=40)
        for m in self.db.fetch("SELECT * FROM masters ORDER BY name"):
            f = ctk.CTkFrame(scroll, fg_color="#1a1a1a", height=105, corner_radius=15)
            f.pack(fill="x", pady=8); f.pack_propagate(False)
            info = ctk.CTkFrame(f, fg_color="transparent")
            info.pack(side="left", padx=30, fill="y")
            ctk.CTkLabel(info, text=m['name'], font=("Arial", 22, "bold")).pack(pady=(10, 0), anchor="w")
            ctk.CTkLabel(info, text=f"Спец: {m.get('specialization') or 'Все'}", font=("Arial", 14), text_color="#555").pack(anchor="w")
            ctk.CTkButton(f, text="РЕДАКТИРОВАТЬ", width=150, height=45, fg_color="#262626",
                          command=lambda d=m: MasterModal(self, self.db, d, self.show_masters)).pack(side="right", padx=30)

    def show_analytics(self):
        self.clear_view()
        hdr = ctk.CTkFrame(self.main_view, fg_color="transparent")
        hdr.pack(fill="x", padx=50, pady=20)
        ctk.CTkLabel(hdr, text="👑 АНАЛИТИКА", font=("Arial", 38, "bold")).pack(side="left")

        filter_f = ctk.CTkFrame(self.main_view, fg_color="#1a1a1a", height=60, corner_radius=10)
        filter_f.pack(fill="x", padx=50, pady=10)
        
        self.dt_start = ctk.CTkEntry(filter_f, width=120)
        self.dt_start.pack(side="left", padx=10)
        self.dt_start.insert(0, (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d"))

        self.dt_end = ctk.CTkEntry(filter_f, width=120)
        self.dt_end.pack(side="left", padx=10)
        self.dt_end.insert(0, datetime.now().strftime("%Y-%m-%d"))

        ctk.CTkButton(filter_f, text="Обновить", command=self.load_analytics_data).pack(side="left", padx=20)
        ctk.CTkButton(filter_f, text="СГЕНЕРИРОВАТЬ ОТЧЕТ (WORD)", fg_color="#2ecc71", 
                      command=self.export_to_word).pack(side="right", padx=20)

        self.stats_container = ctk.CTkFrame(self.main_view, fg_color="transparent")
        self.stats_container.pack(fill="both", expand=True, padx=50, pady=20)
        self.load_analytics_data()

    def load_analytics_data(self):
        for w in self.stats_container.winfo_children(): w.destroy()
        s_date, e_date = f"{self.dt_start.get()} 00:00:00", f"{self.dt_end.get()} 23:59:59"
        rev, cli = self.db.get_stats_summary(s_date, e_date)
        top_m = self.db.get_top_master(s_date, e_date)
        top_c = self.db.get_top_client(s_date, e_date)

        cards = [("ВЫРУЧКА", f"{int(rev):,} ₽", "#2ecc71"), ("КЛИЕНТОВ", str(cli), "#3498db"),
                 ("ТОП МАСТЕР", top_m['name'], "#f1c40f"), ("ТОП КЛИЕНТ", top_c['name'], "#9b59b6")]
        
        for idx, (label, val, color) in enumerate(cards):
            card = ctk.CTkFrame(self.stats_container, fg_color="#1a1a1a", width=300, height=150, corner_radius=20)
            card.grid(row=idx//2, column=idx%2, padx=20, pady=20); card.pack_propagate(False)
            ctk.CTkLabel(card, text=label, font=("Arial", 14, "bold"), text_color="#777").pack(pady=(20, 10))
            ctk.CTkLabel(card, text=str(val), font=("Arial", 28, "bold"), text_color=color).pack()

    def export_to_word(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not path: return
        try:
            doc = Document()
            s_date, e_date = f"{self.dt_start.get()} 00:00:00", f"{self.dt_end.get()} 23:59:59"
            
            doc.add_heading('ОТЧЕТ: КРАСНАЯ КОРОЛЕВА', 0)
            doc.add_paragraph(f"Период отчета: {self.dt_start.get()} — {self.dt_end.get()}")

            # 1. ОБЩАЯ СВОДКА
            doc.add_heading('1. Общие показатели', level=1)
            rev, cli = self.db.get_stats_summary(s_date, e_date)
            table_gen = doc.add_table(rows=3, cols=2)
            table_gen.style = 'Table Grid'
            table_gen.rows[0].cells[0].text = 'Общая выручка'
            table_gen.rows[0].cells[1].text = f"{int(rev)} руб."
            table_gen.rows[1].cells[0].text = 'Всего визитов'
            table_gen.rows[1].cells[1].text = str(cli)
            table_gen.rows[2].cells[0].text = 'Средний чек'
            table_gen.rows[2].cells[1].text = f"{int(rev/cli if cli > 0 else 0)} руб."

            # 2. ВОСТРЕБОВАННОСТЬ УСЛУГ
            doc.add_heading('2. Востребованность услуг', level=1)
            service_data = self.db.get_service_demand_report(s_date, e_date)
            if service_data:
                t_serv = doc.add_table(rows=1, cols=3)
                t_serv.style = 'Table Grid'
                h = t_serv.rows[0].cells
                h[0].text, h[1].text, h[2].text = 'Услуга', 'Записей', 'Заработок'
                for item in service_data:
                    r = t_serv.add_row().cells
                    r[0].text, r[1].text, r[2].text = item['name'], str(item['count']), f"{int(item['total_revenue'])} р."
            
            # 3. ЭФФЕКТИВНОСТЬ МАСТЕРОВ
            doc.add_heading('3. Статистика по мастерам', level=1)
            master_stats = self.db.fetch("""
                SELECT m.name, COUNT(a.id) as count, SUM(s.price) as revenue
                FROM masters m
                JOIN appointments a ON m.id = a.master_id
                JOIN appointment_services aserv ON a.id = aserv.appointment_id
                JOIN services s ON aserv.service_id = s.id
                WHERE a.status = 'Завершено' AND a.appointment_date BETWEEN %s AND %s
                GROUP BY m.name ORDER BY revenue DESC
            """, (s_date, e_date))
            if master_stats:
                t_mast = doc.add_table(rows=1, cols=3)
                t_mast.style = 'Table Grid'
                h = t_mast.rows[0].cells
                h[0].text, h[1].text, h[2].text = 'Мастер', 'Визитов', 'Выручка'
                for m in master_stats:
                    r = t_mast.add_row().cells
                    r[0].text, r[1].text, r[2].text = m['name'], str(m['count']), f"{int(m['revenue'])} р."

            # 4. ТОП КЛИЕНТОВ
            doc.add_heading('4. Топ клиентов по тратам', level=1)
            client_stats = self.db.fetch("""
                SELECT c.name, SUM(s.price) as spent
                FROM clients c
                JOIN appointments a ON c.id = a.client_id
                JOIN appointment_services aserv ON a.id = aserv.appointment_id
                JOIN services s ON aserv.service_id = s.id
                WHERE a.status = 'Завершено' AND a.appointment_date BETWEEN %s AND %s
                GROUP BY c.name ORDER BY spent DESC LIMIT 10
            """, (s_date, e_date))
            if client_stats:
                t_cli = doc.add_table(rows=1, cols=2)
                t_cli.style = 'Table Grid'
                h = t_cli.rows[0].cells
                h[0].text, h[1].text = 'ФИО Клиента', 'Сумма оплат'
                for c in client_stats:
                    r = t_cli.add_row().cells
                    r[0].text, r[1].text = c['name'], f"{int(c['spent'])} р."

            doc.save(path)
            messagebox.showinfo("Успех", f"Подробный аналитический отчет сохранен в {path}")
            os.startfile(path)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать отчет: {e}")

if __name__ == "__main__":
    app = SilkWayApp()
    app.mainloop()